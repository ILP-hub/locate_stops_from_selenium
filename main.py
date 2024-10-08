from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
import time
from docx import Document
import os

# Установите путь к вашему веб-драйверу
driver_path = 'chromedriver.exe'

# Настройки браузера
chrome_options = Options()
chrome_options.add_argument("--new-window")

# Создаем экземпляр веб-драйвера
service = Service(driver_path)
driver = webdriver.Chrome(service=service, options=chrome_options)

# URL страницы с автобусом на Яндексе
bus_urls = [
    '',
    # Добавьте другие URL-адреса по вашему выбору
]

try:
    # Создаем или открываем существующий документ
    file_path = 'bus_routes.docx'
    if os.path.exists(file_path):
        doc = Document(file_path)
    else:
        doc = Document()

    for bus_url in bus_urls:
        # Открываем страницу с автобусом
        driver.get(bus_url)
        time.sleep(3)  # Ждем, чтобы страница полностью загрузилась

        # Ищем название автобуса
        bus_title_element = driver.find_element(By.CLASS_NAME, 'masstransit-card-header-view__title')
        bus_title = bus_title_element.text if bus_title_element else "Название автобуса не найдено"

        # Ищем элемент с классом 'masstransit-legend-group-view__item-link'
        route_links = driver.find_elements(By.CLASS_NAME, 'masstransit-legend-group-view__item-link')

        # Сохраняем ссылки на маршруты
        route_urls = [link.get_attribute('href') for link in route_links]

        # Список для хранения названий остановок и ссылок
        stops_info = []

        # Открываем каждую ссылку на маршрут в новой вкладке и извлекаем информацию
        for route_url in route_urls:
            driver.execute_script("window.open(arguments[0]);", route_url)
            driver.switch_to.window(driver.window_handles[-1])
            time.sleep(3)  # Ждем, чтобы страница полностью загрузилась

            # Ищем названия остановок (примерный селектор, нужно уточнить по HTML страницы)
            stops = driver.find_elements(By.CLASS_NAME, 'card-title-view__title')
            stop_names = [stop.text for stop in stops]

            # Получаем текущий URL из адресной строки
            current_url = driver.current_url

            # Сохраняем информацию
            stops_info.append({
                'route_url': current_url,
                'stop_names': stop_names
            })

            # Закрываем текущую вкладку и возвращаемся на основную
            driver.close()
            driver.switch_to.window(driver.window_handles[0])

        # Добавляем информацию о текущем автобусе в документ
        doc.add_heading(f'Автобус: {bus_title}', level=1)

        # Добавляем собранную информацию о маршрутах в документ
        for info in stops_info:
            doc.add_heading(f"Маршрут: {info['route_url']}", level=2)
            
            # Добавляем полный маршрут в одну сторону
            doc.add_heading('Остановки:', level=3)
            for stop in info['stop_names']:
                doc.add_paragraph(stop)

        # Добавляем маршрут обратно после всех маршрутов в одну сторону
        doc.add_heading('Маршрут обратно:', level=2)
        for info in reversed(stops_info):
            # Добавляем ссылки на маршрут обратно
            doc.add_heading(f"Маршрут: {info['route_url']}", level=3)
            # Добавляем перевернутый маршрут обратно с ссылками
            for stop in info['stop_names']:
                doc.add_paragraph(stop)

    # Сохраняем документ
    doc.save(file_path)

finally:
    # Закрываем браузер
    driver.quit()
